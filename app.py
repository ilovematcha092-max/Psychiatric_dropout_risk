# app.py — Psychiatric Dropout Risk (Expert-feedback rev)
# - Monotone XGBoost + (optional) isotonic calibration
# - Pre/Post planning (avoid leakage on followups)
# - Policy Overlay updated per expert comments
# - Consistency guard for Suicidal/Self-harm chief complaint vs flags
# - 30-day followups definition unified (UI + batch)
# - LOS distribution localized (TW acute psych ~3–4 weeks) + Dx-linked
# - Added social factors & prior-dropout rehospitalization
# - SHAP / Ablation / Decision curve / Capacity / Fairness
# - Vignettes template updated to include new fields

import os, re, math, warnings
warnings.filterwarnings("ignore")

import streamlit as st
import pandas as pd
import numpy as np
import joblib
import shap
import matplotlib.pyplot as plt
from io import BytesIO

try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

st.set_page_config(page_title="Psychiatric Dropout Risk", layout="wide")
st.title("🧠 Psychiatric Dropout Risk Predictor (Expert-feedback rev)")

# ==== Math helpers ====
def _sigmoid(x): return 1.0 / (1.0 + np.exp(-x))
def _logit(p, eps=1e-6):
    p = float(np.clip(p, eps, 1 - eps)); return np.log(p / (1 - p))
def _logit_vec(p, eps=1e-6):
    p = np.clip(p, eps, 1 - eps); return np.log(p / (1 - p))

# ==== Global knobs ====
CAL_LOGIT_SHIFT = float(os.getenv("RISK_CAL_SHIFT", "0.0"))
BORDER_BAND = 7
BLEND_W_DEFAULT = 0.30
SOFT_UPLIFT = {"floor": 0.60, "add": 0.10, "cap": 0.90}  # self-harm safety uplift

# Overlay safety controls
OVERLAY_SCALE = 0.75
DELTA_CLIP   = 1.00
TEMP         = 1.20

# ====== Options & schema ======
DIAG_LIST = [
    "Schizophrenia","Bipolar","Depression","Personality Disorder",
    "Substance Use Disorder","Dementia","Anxiety","PTSD","OCD","ADHD","Other/Unknown"
]
BIN_YESNO = ["Yes","No"]
GENDER_LIST = ["Male","Female","Other/Unknown"]

CHIEF_LIST = [
    "Suicidal ideation/attempt","Self-harm (non-suicidal)","Aggression/violence",
    "Severe agitation/mania","Psychosis/diagnostic workup",
    "Poor self-care/failure to thrive","Medication side effects/intoxication","Other/Unknown"
]
BIPOLAR_EP_LIST = ["N/A","Manic","Depressive","Mixed","Hypomanic"]

# unified: 30-day followups
FOLLOWUPS_LABEL = "Post-discharge Followups (30d count)"

# ====== Feature schema (one-hots included) ======
TEMPLATE_COLUMNS = [
    "age","length_of_stay","num_previous_admissions",
    "medication_compliance_score","family_support_score","financial_strain_score",
    "post_discharge_followups_30d",
    "living_alone_Yes","living_alone_No",
    "has_case_manager_Yes","has_case_manager_No",
    "prior_dropout_rehosp_Yes","prior_dropout_rehosp_No",
    "gender_Male","gender_Female","gender_Other/Unknown",
] + [f"diagnosis_{d}" for d in DIAG_LIST] + \
    [f"chief_{c}" for c in CHIEF_LIST] + \
    [f"bipolar_ep_{b}" for b in BIPOLAR_EP_LIST] + [
    "has_recent_self_harm_Yes","has_recent_self_harm_No",
    "self_harm_during_admission_Yes","self_harm_during_admission_No",
]

# ====== Null-safe defaults ======
DEFAULTS = {
    "age": 40.0,
    "length_of_stay": 21.0,  # closer to TW acute psych ward avg 3–4 weeks
    "num_previous_admissions": 0.0,
    "medication_compliance_score": 5.0,
    "family_support_score": 5.0,
    "financial_strain_score": 5.0,
    "post_discharge_followups_30d": 0.0,
}
NUMERIC_KEYS = list(DEFAULTS.keys())

def _num_or_default(x, key):
    try: v = float(x)
    except Exception: v = np.nan
    if pd.isna(v): v = DEFAULTS.get(key, 0.0)
    return v

def fill_defaults_single_row(X1: pd.DataFrame):
    i = X1.index[0]
    for k in NUMERIC_KEYS:
        if k in X1.columns:
            X1.at[i, k] = _num_or_default(X1.at[i, k], k)
    # one-hot fallbacks
    diag_cols = [c for c in X1.columns if c.startswith("diagnosis_")]
    if diag_cols and (X1.loc[i, diag_cols].sum() == 0):
        col = "diagnosis_Other/Unknown"
        if col in X1.columns: X1.at[i, col] = 1
    chief_cols = [c for c in X1.columns if c.startswith("chief_")]
    if chief_cols and (X1.loc[i, chief_cols].sum() == 0):
        X1.at[i, "chief_Other/Unknown"] = 1
    bip_cols = [c for c in X1.columns if c.startswith("bipolar_ep_")]
    if bip_cols and (X1.loc[i, bip_cols].sum() == 0):
        X1.at[i, "bipolar_ep_N/A"] = 1
    oh_cols = [c for c in X1.columns if "_" in c and c not in NUMERIC_KEYS]
    if oh_cols: X1.loc[i, oh_cols] = X1.loc[i, oh_cols].fillna(0)

def fill_defaults_batch(df_feat: pd.DataFrame):
    for k in NUMERIC_KEYS:
        if k in df_feat.columns:
            df_feat[k] = pd.to_numeric(df_feat[k], errors="coerce").fillna(DEFAULTS[k])
    oh_cols = [c for c in df_feat.columns if "_" in c and c not in NUMERIC_KEYS]
    if oh_cols: df_feat[oh_cols] = df_feat[oh_cols].fillna(0)
    # fallbacks
    diag_cols = [c for c in df_feat.columns if c.startswith("diagnosis_")]
    if diag_cols:
        none_diag = (df_feat[diag_cols].sum(axis=1) == 0)
        if none_diag.any() and "diagnosis_Other/Unknown" in df_feat.columns:
            df_feat.loc[none_diag, "diagnosis_Other/Unknown"] = 1
    chief_cols = [c for c in df_feat.columns if c.startswith("chief_")]
    if chief_cols:
        none_chief = (df_feat[chief_cols].sum(axis=1) == 0)
        if none_chief.any() and "chief_Other/Unknown" in df_feat.columns:
            df_feat.loc[none_chief, "chief_Other/Unknown"] = 1
    bip_cols = [c for c in df_feat.columns if c.startswith("bipolar_ep_")]
    if bip_cols:
        none_bip = (df_feat[bip_cols].sum(axis=1) == 0)
        if none_bip.any() and "bipolar_ep_N/A" in df_feat.columns:
            df_feat.loc[none_bip, "bipolar_ep_N/A"] = 1

# ====== Overlay policy（log-odds, directionally monotone）======
POLICY = {
    # clinical history / social
    "per_prev_admission": 0.18,
    "per_point_low_support": 0.18,
    "per_point_financial_strain": 0.12,
    "per_followup": -0.18,              # more followups (30d) protective
    "no_followup_extra": 0.30,
    "living_alone": 0.20,
    "case_manager_protect": -0.25,
    "prior_dropout_rehosp": 0.40,       # strong penalty

    # LOS
    "los_short": 0.30, "los_mid": 0.00, "los_mid_high": 0.10, "los_long": 0.20,

    # age ends
    "age_young": 0.10, "age_old": 0.10,

    # Dx weights
    "diag": {
        "Personality Disorder":    0.35,
        "Substance Use Disorder":  0.35,
        "Bipolar":                 0.12,
        "PTSD":                    0.12,
        "Schizophrenia":           0.12,
        "Depression":              0.05,
        "Anxiety":                 0.00,
        "OCD":                     0.00,
        "Dementia":                0.00,
        "ADHD":                    0.00,
        "Other/Unknown":           0.00,
    },

    # Bipolar episode subtype
    "bip_ep": {
        "Manic": 0.20, "Depressive": 0.10, "Mixed": 0.25, "Hypomanic": 0.05, "N/A": 0.00
    },

    # Chief complaint
    "chief": {
        "Suicidal ideation/attempt": 0.50,
        "Self-harm (non-suicidal)":  0.35,
        "Aggression/violence":       0.25,
        "Severe agitation/mania":    0.20,
        "Psychosis/diagnostic workup": 0.15,
        "Poor self-care/failure to thrive": 0.10,
        "Medication side effects/intoxication": 0.05,
        "Other/Unknown": 0.00
    },

    # interactions
    "x_sud_lowcomp": 0.30,
    "x_pd_shortlos": 0.10,

    # compliance (center=5)
    "per_point_low_compliance": 0.22,
    "per_point_high_compliance_protect": -0.06,
}

# ====== UI — Sidebar ======
with st.sidebar:
    st.header("Patient Info")
    age = st.slider("Age", 18, 95, 35)
    gender = st.selectbox("Gender", GENDER_LIST, index=0)
    diagnoses = st.multiselect("Diagnoses (multi-select)", DIAG_LIST, default=[])
    chief = st.multiselect("Chief Complaint(s)", CHIEF_LIST, default=[])
    bip_ep = st.selectbox("Bipolar episode subtype", BIPOLAR_EP_LIST, index=0)

    length_of_stay = st.slider("Length of Stay (days)", 1, 90, 21)
    num_adm = st.slider("Previous Admissions (1y)", 0, 15, 1)

    compliance = st.slider("Medication Compliance (0–10)", 0.0, 10.0, 5.0, 0.5)
    support = st.slider("Family Support (0–10)", 0.0, 10.0, 5.0, 0.5)
    financial = st.slider("Financial Strain (0–10)", 0.0, 10.0, 5.0, 0.5)

    followups_30d = st.slider(FOLLOWUPS_LABEL, 0, 12, 2)

    living_alone = st.radio("Living alone", BIN_YESNO, index=1)
    has_cm = st.radio("Has case manager", BIN_YESNO, index=1)
    prior_drop_reh = st.radio("Prior dropout→rehospitalization (1y)", BIN_YESNO, index=1)

    recent_self_harm = st.radio("Recent Self-harm", BIN_YESNO, index=1)
    selfharm_adm = st.radio("Self-harm During Admission", BIN_YESNO, index=1)

    mode = st.radio("Mode", ["Pre-planning (no followups feature)", "Post-planning (monitoring)"], index=0)
    use_followups_feature = (mode.startswith("Post"))

    st.markdown("---")
    with st.expander("Advanced (calibration & overlay)", expanded=False):
        cal_shift = st.slider("Global calibration (log-odds shift)", -1.0, 1.0, CAL_LOGIT_SHIFT, 0.05)
        blend_w  = st.slider("Blend weight (Final = (1-BLEND)*Model + BLEND*Overlay)", 0.0, 1.0, BLEND_W_DEFAULT, 0.05)
        overlay_scale = st.slider("Overlay scale", 0.0, 1.0, OVERLAY_SCALE, 0.05)
        delta_clip = st.slider("Overlay delta clip |log-odds|", 0.0, 2.0, DELTA_CLIP, 0.05)
        temp_val = st.slider("Temperature (>1 softer probs)", 0.5, 3.0, TEMP, 0.05)
    CAL_LOGIT_SHIFT, BLEND_W, OVERLAY_SCALE, DELTA_CLIP, TEMP = cal_shift, blend_w, overlay_scale, delta_clip, temp_val

# ====== Helpers (one-hot, alignment) ======
def set_onehot_by_prefix(df, prefix, value):
    col = f"{prefix}_{value}"
    if col in df.columns: df.at[0, col] = 1

def set_onehot_by_prefix_multi(df, prefix, values):
    for v in values:
        col = f"{prefix}_{v}"
        if col in df.columns: df.at[0, col] = 1

def flag_yes(row, prefix):
    col = f"{prefix}_Yes"; return (col in row.index) and (row[col] == 1)

def risk_bins(score, mod=20, high=40, band=BORDER_BAND):
    if score >= high + band: return "High"
    if score >= high - band: return "Moderate–High"
    if score >= mod + band:  return "Moderate"
    if score >= mod - band:  return "Low–Moderate"
    return "Low"

def proba_to_percent(p): return float(p) * 100
def proba_to_score(p): return int(round(proba_to_percent(p)))

# ====== Model load / train (monotone) + optional isotonic calibration ======
def get_monotone_constraints(feature_names):
    mono_map = {
        "num_previous_admissions": +1,
        "medication_compliance_score": -1,
        "family_support_score": -1,
        "financial_strain_score": +1,
        "post_discharge_followups_30d": -1,
        "length_of_stay": +1,
        "living_alone_Yes": +1,
        "has_case_manager_Yes": -1,
        "prior_dropout_rehosp_Yes": +1,
        "has_recent_self_harm_Yes": +1,
        "self_harm_during_admission_Yes": +1,
    }
    cons = [str(mono_map.get(f, 0)) for f in feature_names]
    return "(" + ",".join(cons) + ")"

def xgb_model_with_monotone(feature_names):
    import xgboost as xgb
    return xgb.XGBClassifier(
        n_estimators=500, max_depth=4, learning_rate=0.06,
        subsample=0.9, colsample_bytree=0.9, reg_lambda=1.0,
        random_state=42, tree_method="hist",
        objective="binary:logistic", eval_metric="logloss",
        monotone_constraints=get_monotone_constraints(feature_names)
    )

def try_load_model(path="dropout_model.pkl"):
    if not os.path.exists(path): return None
    try:
        mdl = joblib.load(path)
        if isinstance(mdl, dict) and "model" in mdl: return mdl["model"]
        return mdl
    except Exception:
        return None

def align_df_to_model(df: pd.DataFrame, m):
    names = None
    try:
        booster = getattr(m, "get_booster", lambda: None)()
        if booster is not None:
            nm = getattr(booster, "feature_names", None)
            if nm: names = list(nm)
    except Exception: pass
    if names:
        aligned = pd.DataFrame(0, index=df.index, columns=names, dtype=np.float32)
        inter = [c for c in names if c in df.columns]
        aligned.loc[:, inter] = df[inter].astype(np.float32).values
        return aligned, names
    out = df.astype(np.float32)
    return out, list(out.columns)

def train_demo_model_and_calibrator(columns):
    import xgboost as xgb
    from sklearn.model_selection import train_test_split
    from sklearn.calibration import CalibratedClassifierCV

    rng = np.random.default_rng(42)
    n = 12000
    X = pd.DataFrame(0, index=range(n), columns=columns, dtype=np.float32)

    # base numeric
    X["age"] = rng.integers(16, 85, n)
    # LOS ~ N(22, 8), clip 1–60, and add Dx-linked shifts
    base_los = rng.normal(22.0, 8.0, n)
    base_los = np.clip(base_los, 1, 60)
    X["length_of_stay"] = base_los
    X["num_previous_admissions"] = rng.poisson(1.0, n).clip(0, 12)
    X["medication_compliance_score"] = rng.normal(6.0, 2.5, n).clip(0, 10)
    X["family_support_score"] = rng.normal(5.0, 2.5, n).clip(0, 10)
    X["financial_strain_score"] = rng.normal(5.0, 2.5, n).clip(0, 10)
    X["post_discharge_followups_30d"] = rng.integers(0, 6, n)

    # gender
    idx_gender = rng.integers(0, len(GENDER_LIST), n)
    for i, g in enumerate(GENDER_LIST): X.loc[idx_gender == i, f"gender_{g}"] = 1

    # Dx + comorb
    idx_primary = rng.integers(0, len(DIAG_LIST), n)
    for i, d in enumerate(DIAG_LIST): X.loc[idx_primary == i, f"diagnosis_{d}"] = 1
    for d, pr in {"Substance Use Disorder": 0.22, "Depression": 0.27, "Anxiety": 0.22, "PTSD": 0.12}.items():
        X.loc[rng.random(n) < pr, f"diagnosis_{d}"] = 1

    # LOS shift by Dx
    los = X["length_of_stay"].to_numpy()
    los += 5.0 * (X["diagnosis_Schizophrenia"]==1).to_numpy()
    los += 3.0 * (X["diagnosis_Bipolar"]==1).to_numpy()
    X["length_of_stay"] = np.clip(los, 1, 90)

    # Chief complaint
    for c in CHIEF_LIST:
        X.loc[rng.random(n) < 0.15, f"chief_{c}"] = 1
    # enforce at least one
    chief_cols = [f"chief_{c}" for c in CHIEF_LIST]
    none_chief = (X[chief_cols].sum(axis=1) == 0)
    X.loc[none_chief, "chief_Other/Unknown"] = 1

    # Bipolar episode subtype
    probs_bip = [0.70, 0.07, 0.10, 0.05, 0.08]  # mostly N/A
    choices = rng.choice(len(BIPOLAR_EP_LIST), size=n, p=probs_bip)
    for i, b in enumerate(BIPOLAR_EP_LIST):
        X.loc[choices == i, f"bipolar_ep_{b}"] = 1

    # social flags
    X["living_alone_Yes"] = (rng.random(n) < 0.25).astype(int); X["living_alone_No"] = 1 - X["living_alone_Yes"]
    X["has_case_manager_Yes"] = (rng.random(n) < 0.35).astype(int); X["has_case_manager_No"] = 1 - X["has_case_manager_Yes"]
    X["prior_dropout_rehosp_Yes"] = (rng.random(n) < 0.12).astype(int); X["prior_dropout_rehosp_No"] = 1 - X["prior_dropout_rehosp_Yes"]

    # self-harm flags
    r1, r2 = rng.integers(0, 2, n), rng.integers(0, 2, n)
    X.loc[r1 == 1, "has_recent_self_harm_Yes"] = 1; X.loc[r1 == 0, "has_recent_self_harm_No"] = 1
    X.loc[r2 == 1, "self_harm_during_admission_Yes"] = 1; X.loc[r2 == 0, "self_harm_during_admission_No"] = 1

    # target logit (directionally consistent)
    beta0 = -0.50
    prev_ge2 = (X["num_previous_admissions"] >= 2).to_numpy().astype(np.float32)

    def col(name): return X[name].to_numpy().astype(np.float32)

    logit = (beta0
        + 0.80*col("has_recent_self_harm_Yes")
        + 0.60*col("self_harm_during_admission_Yes")
        + 0.60*prev_ge2
        - 0.25*col("medication_compliance_score")
        - 0.20*col("family_support_score")
        + 0.12*col("financial_strain_score")
        - 0.15*col("post_discharge_followups_30d")
        + 0.05*col("length_of_stay"))

    for dx, w in POLICY["diag"].items():
        nm=f"diagnosis_{dx}"
        if nm in X.columns: logit += w*col(nm)

    for ep, w in POLICY["bip_ep"].items():
        nm=f"bipolar_ep_{ep}"
        if nm in X.columns: logit += w*col(nm)

    for cf, w in POLICY["chief"].items():
        nm=f"chief_{cf}"
        if nm in X.columns: logit += w*col(nm)

    logit += 0.20*col("living_alone_Yes") - 0.25*col("has_case_manager_Yes") + 0.40*col("prior_dropout_rehosp_Yes")

    # interactions
    if "diagnosis_Substance Use Disorder" in X.columns:
        logit += 0.30*((col("diagnosis_Substance Use Disorder")==1) & (col("medication_compliance_score")<=3))
    if "diagnosis_Personality Disorder" in X.columns:
        logit += 0.10*((col("diagnosis_Personality Disorder")==1) & (col("length_of_stay")<3))

    noise = rng.normal(0.0, 0.35, n).astype(np.float32)
    p = 1.0 / (1.0 + np.exp(-(logit + noise)))
    y = (rng.random(n) < p).astype(np.int32)

    model = xgb_model_with_monotone(list(X.columns))
    model.fit(X, y)

    calibrator = None
    try:
        X_tr, X_ca, y_tr, y_ca = train_test_split(X, y, test_size=0.2, random_state=777)
        calibrator = CalibratedClassifierCV(model, cv="prefit", method="isotonic")
        calibrator.fit(X_ca, y_ca)
    except Exception:
        calibrator = None

    return model, calibrator, "demo (monotone + isotonic if available)"

def get_feat_names(m):
    try:
        b = m.get_booster()
        if getattr(b, "feature_names", None): return list(b.feature_names)
    except Exception: pass
    if hasattr(m, "feature_names_in_"): return list(m.feature_names_in_)
    return None

_loaded = try_load_model()
if _loaded is None:
    model, calibrator, model_source = train_demo_model_and_calibrator(TEMPLATE_COLUMNS)
else:
    model, calibrator, model_source = _loaded, None, "loaded from dropout_model.pkl"

def predict_model_proba(df_aligned: pd.DataFrame):
    probs = model.predict_proba(df_aligned, validate_features=False)[:, 1]
    if calibrator is not None:
        try:
            probs = calibrator.predict_proba(df_aligned)[:, 1]
        except Exception:
            pass
    return probs

# ====== Overlay（單例 + 驅動因子）======
def overlay_single_and_drivers(X1: pd.DataFrame, base_prob: float, include_followup_effect: bool = True):
    row = X1.iloc[0]
    drivers = []
    def add(label, val):
        if val != 0: drivers.append((label, float(val)))
        return val
    base_logit = _logit(float(base_prob))
    lz = base_logit

    # pull values
    adm = _num_or_default(row["num_previous_admissions"], "num_previous_admissions")
    comp = _num_or_default(row["medication_compliance_score"], "medication_compliance_score")
    sup  = _num_or_default(row["family_support_score"], "family_support_score")
    fin  = _num_or_default(row["financial_strain_score"], "financial_strain_score")
    fup  = _num_or_default(row["post_discharge_followups_30d"], "post_discharge_followups_30d")
    los  = _num_or_default(row["length_of_stay"], "length_of_stay")
    agev = _num_or_default(row["age"], "age")

    lz += add("More previous admissions", POLICY["per_prev_admission"] * min(int(adm), 5))
    lz += add("Low family support", POLICY["per_point_low_support"] * max(0.0, 5.0 - sup))
    lz += add("Financial strain", POLICY["per_point_financial_strain"] * max(0.0, fin - 5.0))

    # compliance：低→加分，高→保護
    lz += add("Low medication compliance", POLICY["per_point_low_compliance"] * max(0.0, 5.0 - comp))
    if comp >= 8:
        lz += add("High compliance (protective)", POLICY["per_point_high_compliance_protect"] * (comp - 7.0))

    if include_followup_effect:
        lz += add("More followups in 30d (protective)", POLICY["per_followup"] * fup)
        if fup == 0: lz += add("No follow-up in 30d", POLICY["no_followup_extra"])

    if los < 3: lz += add("Very short stay (<3d)", POLICY["los_short"])
    elif los <= 21: lz += add("Typical stay (3–21d)", POLICY["los_mid"])
    elif los <= 28: lz += add("Longish stay (22–28d)", POLICY["los_mid_high"])
    else: lz += add("Very long stay (>28d)", POLICY["los_long"])

    if agev < 21: lz += add("Young age (<21)", POLICY["age_young"])
    elif agev >= 75: lz += add("Older age (≥75)", POLICY["age_old"])

    # social flags
    if row.get("living_alone_Yes", 0) == 1: lz += add("Living alone", POLICY["living_alone"])
    if row.get("has_case_manager_Yes", 0) == 1: lz += add("Has case manager (protective)", POLICY["case_manager_protect"])
    if row.get("prior_dropout_rehosp_Yes", 0) == 1: lz += add("Prior dropout→rehosp", POLICY["prior_dropout_rehosp"])

    # Dx
    for dx, w in POLICY["diag"].items():
        key = f"diagnosis_{dx}"
        if X1.at[0, key] == 1: lz += add(f"Diagnosis: {dx}", w)

    # Bipolar episode
    for ep, w in POLICY["bip_ep"].items():
        key = f"bipolar_ep_{ep}"
        if key in X1.columns and X1.at[0, key] == 1: lz += add(f"Bipolar episode: {ep}", w)

    # Chief complaints
    for cf, w in POLICY["chief"].items():
        key = f"chief_{cf}"
        if key in X1.columns and X1.at[0, key] == 1: lz += add(f"Chief: {cf}", w)

    # interactions
    if (X1.at[0, "diagnosis_Substance Use Disorder"] == 1) and (comp <= 3):
        lz += add("SUD × very low compliance", POLICY["x_sud_lowcomp"])
    if (X1.at[0, "diagnosis_Personality Disorder"] == 1) and (los < 3):
        lz += add("PD × very short stay", POLICY["x_pd_shortlos"])

    # scale + clip + calibration + temp
    delta = np.clip(OVERLAY_SCALE * (lz - base_logit), -DELTA_CLIP, DELTA_CLIP)
    lz2 = base_logit + delta + CAL_LOGIT_SHIFT
    p_overlay = _sigmoid(lz2 / TEMP)
    return float(p_overlay), drivers

# ====== Build single-row DF ======
X_single = pd.DataFrame(0, index=[0], columns=TEMPLATE_COLUMNS, dtype=float)
for k, v in {
    "age": age, "length_of_stay": float(length_of_stay), "num_previous_admissions": int(num_adm),
    "medication_compliance_score": float(compliance), "family_support_score": float(support),
    "financial_strain_score": float(financial),
    "post_discharge_followups_30d": int(followups_30d),
}.items(): X_single.at[0, k] = v
set_onehot_by_prefix(X_single, "gender", gender)
set_onehot_by_prefix_multi(X_single, "diagnosis", diagnoses)
set_onehot_by_prefix_multi(X_single, "chief", chief)
set_onehot_by_prefix(X_single, "bipolar_ep", bip_ep)
set_onehot_by_prefix(X_single, "living_alone", living_alone)
set_onehot_by_prefix(X_single, "has_case_manager", has_cm)
set_onehot_by_prefix(X_single, "prior_dropout_rehosp", prior_drop_reh)
set_onehot_by_prefix(X_single, "has_recent_self_harm", recent_self_harm)
set_onehot_by_prefix(X_single, "self_harm_during_admission", selfharm_adm)
fill_defaults_single_row(X_single)

# Consistency guard: suicidal/self-harm chief but flags==No -> auto treat as risk for computation
need_selfharm_patch = False
if (("chief_Suicidal ideation/attempt" in X_single.columns and X_single.at[0,"chief_Suicidal ideation/attempt"]==1) or
    ("chief_Self-harm (non-suicidal)" in X_single.columns and X_single.at[0,"chief_Self-harm (non-suicidal)"]==1)):
    if (X_single.at[0,"has_recent_self_harm_Yes"]==0) and (X_single.at[0,"self_harm_during_admission_Yes"]==0):
        need_selfharm_patch = True

X_used = X_single.copy()
if need_selfharm_patch:
    X_used.at[0,"has_recent_self_harm_Yes"] = 1  # for computation
    st.info("⚠️ Chief complaint indicates suicidality/self-harm, but self-harm flags were 'No'. "
            "Temporarily treating as self-harm risk for prediction. Please confirm the flags.")

# Pre-planning：避免洩漏 → followups=0
if not use_followups_feature:
    X_used.at[0, "post_discharge_followups_30d"] = 0

# 預測（model + overlay + blend + uplift）
X_align, _ = align_df_to_model(X_used, model)
p_model = float(predict_model_proba(X_align)[0])
p_overlay, drivers = overlay_single_and_drivers(X_used, base_prob=p_model, include_followup_effect=use_followups_feature)
p_final = (1.0 - BLEND_W) * p_model + BLEND_W * p_overlay

# 自傷 uplift
if flag_yes(X_used.iloc[0], "has_recent_self_harm") or flag_yes(X_used.iloc[0], "self_harm_during_admission"):
    p_final = min(max(p_final, SOFT_UPLIFT["floor"]) + SOFT_UPLIFT["add"], SOFT_UPLIFT["cap"])

percent_model = proba_to_percent(p_model)
percent_overlay = proba_to_percent(p_overlay)
percent_final = proba_to_percent(p_final)
score = proba_to_score(p_final)
level = risk_bins(score)

# ====== Guards / input reasonableness ======
warns = []
if length_of_stay < 3: warns.append("Very short LOS (<3d) is uncommon in TW acute psych; confirm context.")
if length_of_stay > 60: warns.append("LOS > 60d is unusual for acute; check definition.")
if num_adm > 10: warns.append("Previous admissions >10 in 1y is uncommon; confirm definition.")
if warns: st.info("ℹ️ Data sanity check:\n- " + "\n- ".join(warns))

# ====== Show results ======
st.subheader("Predicted Dropout Risk (within 3 months)")
c1, c2, c3, c4 = st.columns(4)
with c1: st.metric("Model Probability", f"{percent_model:.1f}%")
with c2: st.metric("Overlay Probability", f"{percent_overlay:.1f}%")
with c3: st.metric("Final Probability", f"{percent_final:.1f}%")
with c4: st.metric("Risk Score (0–100)", f"{score}")

note = "Pre-planning mode: followups feature is ignored." if not use_followups_feature else "Post-planning mode: followups feature is used (30-day count)."
st.caption(note)

if level == "High":
    st.error("🔴 High Risk")
elif level == "Moderate–High":
    st.warning("🟠 Moderate–High (borderline to High)")
elif level == "Moderate":
    st.warning("🟡 Moderate Risk")
elif level == "Low–Moderate":
    st.info("🔵 Low–Moderate (borderline to Moderate)")
else:
    st.success("🟢 Low Risk")

# ====== SHAP + Policy drivers + 對齊卡 ======
with st.expander("🔍 Explanations — Model SHAP vs Policy drivers", expanded=True):
    import xgboost as xgb
    try:
        booster = model.get_booster()
        dmat = xgb.DMatrix(X_align, feature_names=list(X_align.columns))
        contribs = booster.predict(dmat, pred_contribs=True, validate_features=False)
        contrib = np.asarray(contribs)[0]
        base_value = float(contrib[-1])
        sv_map = dict(zip(list(X_align.columns), contrib[:-1]))
    except Exception:
        explainer = shap.TreeExplainer(model)
        sv_raw = explainer.shap_values(X_align)
        base_value = explainer.expected_value
        if isinstance(base_value, (list, np.ndarray)) and not np.isscalar(base_value):
            base_value = base_value[0]
            if isinstance(sv_raw, list): sv_raw = sv_raw[0]
        sv_map = dict(zip(list(X_align.columns), np.array(sv_raw)[0]))

    feat_rows = []
    def _push_shap(label, key, shown_value):
        if key in sv_map:
            feat_rows.append({"feature": label, "value": shown_value, "model_shap": float(sv_map[key]), "key": key})

    # numeric
    _push_shap("Age", "age", X_used.at[0,"age"])
    _push_shap("Length of Stay", "length_of_stay", X_used.at[0,"length_of_stay"])
    _push_shap("Previous Admissions", "num_previous_admissions", X_used.at[0,"num_previous_admissions"])
    _push_shap("Medication Compliance", "medication_compliance_score", X_used.at[0,"medication_compliance_score"])
    _push_shap("Family Support", "family_support_score", X_used.at[0,"family_support_score"])
    _push_shap("Financial Strain", "financial_strain_score", X_used.at[0,"financial_strain_score"])
    _push_shap("Followups (30d)", "post_discharge_followups_30d", X_used.at[0,"post_discharge_followups_30d"])
    # one-hots picked
    for dx in diagnoses: _push_shap(f"Diagnosis={dx}", f"diagnosis_{dx}", 1)
    for cf in chief: _push_shap(f"Chief={cf}", f"chief_{cf}", 1)
    _push_shap(f"Bipolar episode={bip_ep}", f"bipolar_ep_{bip_ep}", 1)
    _push_shap(f"Gender={gender}", f"gender_{gender}", 1)
    _push_shap(f"Living alone={living_alone}", f"living_alone_{living_alone}", 1)
    _push_shap(f"Has case manager={has_cm}", f"has_case_manager_{has_cm}", 1)
    _push_shap(f"Prior dropout→rehosp={prior_drop_reh}", f"prior_dropout_rehosp_{prior_drop_reh}", 1)
    _push_shap(f"Recent Self-harm={recent_self_harm}", f"has_recent_self_harm_{recent_self_harm}", 1)
    _push_shap(f"Self-harm During Admission={selfharm_adm}", f"self_harm_during_admission_{selfharm_adm}", 1)

    df_shap = pd.DataFrame(feat_rows)

    if len(df_shap):
        df_top = df_shap.reindex(df_shap["model_shap"].abs().sort_values(ascending=False).index).head(12)
        names = df_top["feature"].tolist()
        vals = df_top["model_shap"].to_numpy(dtype=float)
        data_vals = df_top["value"].to_numpy(dtype=float)
        exp = shap.Explanation(values=vals, base_values=base_value, feature_names=names, data=data_vals)
        shap.plots.waterfall(exp, show=False, max_display=12)
        st.pyplot(plt.gcf(), clear_figure=True)
        st.caption("Model SHAP (top by |value|)")
        st.dataframe(df_top[["feature","value","model_shap"]], use_container_width=True)
    else:
        st.caption("No SHAP contributions available for the selected case.")

    # Policy drivers
    df_drv = pd.DataFrame(
        [{"driver": k, "policy_log_odds (pre-scale)": round(v, 3)} for k, v in sorted(drivers, key=lambda x: abs(x[1]), reverse=True)]
    )
    st.caption("Policy drivers")
    if len(df_drv): st.dataframe(df_drv, use_container_width=True)

    # Alignment check
    st.caption("Alignment check (Model vs Policy) — look for ⚠️ if directions disagree.")
    def _sign(x): return 1 if x>1e-6 else (-1 if x<-1e-6 else 0)
    align_rows = []
    name_map = [
        ("Previous Admissions","num_previous_admissions","More previous admissions"),
        ("Medication Compliance","medication_compliance_score","Low medication compliance"),
        ("Family Support","family_support_score","Low family support"),
        ("Financial Strain","financial_strain_score","Financial strain"),
        ("Followups (30d)","post_discharge_followups_30d","More followups in 30d (protective)"),
        ("Length of Stay","length_of_stay","Very short stay (<3d) / Long stay"),
    ]
    for lab, key, dname in name_map:
        shap_v = float(sv_map.get(key, 0.0))
        pol = 0.0
        for nm, v in drivers:
            if dname.split()[0] in nm: pol += v
        ms, ps = _sign(shap_v), _sign(pol)
        align_rows.append({"feature": lab, "model_sign": ms, "policy_sign": ps, "flag": "⚠️" if (ms*ps==-1) else ""})
    st.dataframe(pd.DataFrame(align_rows), use_container_width=True)

# ====== Recommended actions ======
st.subheader("Recommended Actions")
BASE_ACTIONS = {
    "High": [
        ("Today","Clinician","Crisis/safety planning; 24/7 crisis contacts"),
        ("Today","Clinic scheduler","Return within 7 days (prefer 72h)"),
        ("Today","Care coordinator","Warm handoff to case management"),
        ("48h","Nurse","Outreach call: symptoms/side-effects/barriers"),
        ("7d","Pharmacist/Nurse","Medication review + adherence aids"),
        ("1–2w","Social worker","SDOH screen; transport/financial aid"),
        ("1–4w","Peer support","Enroll in peer/skills group"),
    ],
    "Moderate": [
        ("1–2w","Clinic scheduler","Book within 14 days; SMS reminders"),
        ("1–2w","Nurse","Barrier check & solutions"),
        ("2–4w","Clinician","Brief MI/BA/psychoeducation; 4-week plan"),
    ],
    "Low": [
        ("2–4w","Clinic scheduler","Routine follow-up; confirm reminders"),
        ("2–4w","Nurse","Education/self-management resources"),
    ],
}
def _normalize_action_tuple(a):
    if len(a)==3: return (a[0],a[1],a[2])
    return a

def personalized_actions(row: pd.Series, chosen_dx: list, chief_list: list):
    acts = []
    comp = _num_or_default(row["medication_compliance_score"], "medication_compliance_score")
    sup  = _num_or_default(row["family_support_score"], "family_support_score")
    fup  = _num_or_default(row["post_discharge_followups_30d"], "post_discharge_followups_30d")
    los  = _num_or_default(row["length_of_stay"], "length_of_stay")
    agev = _num_or_default(row["age"], "age")
    has_selfharm = flag_yes(row, "has_recent_self_harm") or flag_yes(row, "self_harm_during_admission") or \
                   ("Suicidal ideation/attempt" in chief_list) or ("Self-harm (non-suicidal)" in chief_list)
    has_sud = "Substance Use Disorder" in chosen_dx
    has_pd  = "Personality Disorder" in chosen_dx
    has_dep = "Depression" in chosen_dx
    has_scz = "Schizophrenia" in chosen_dx

    if has_selfharm:
        acts += [("Today","Clinician","C-SSRS; update safety plan; lethal-means counseling"),
                 ("48h","Nurse","Safety check-in call")]
    if has_sud and comp <= 3:
        acts += [("1–7d","Clinician","Brief MI focused on use goals"),
                 ("1–7d","Care coordinator","Refer to SUD program/IOP or CM"),
                 ("Today","Clinician","Overdose prevention education")]
    if has_pd and los < 3:
        acts += [("Today","Care coordinator","Same-day DBT/skills intake"),
                 ("48h","Peer support","Proactive outreach + skills workbook")]
    if comp <= 3:
        acts += [("7d","Pharmacist","Simplify regimen + blister/pillbox + reminders"),
                 ("1–2w","Clinician","Consider LAI if appropriate")]
    if sup <= 2:
        acts += [("1–2w","Clinician","Family meeting / caregiver engagement"),
                 ("1–2w","Social worker","Community supports; transport/financial counseling")]
    if fup == 0:
        acts += [("Today","Clinic scheduler","Book 2 touchpoints in first 14 days (day2/day7)")]
    if los < 3:
        acts += [("48h","Nurse","Early call; review meds/barriers")]
    elif los > 28:
        acts += [("1–7d","Care coordinator","Step-down/day program plan + warm handoff")]
    if agev < 21:
        acts += [("1–2w","Clinician","Involve guardians; link school counseling")]
    elif agev >= 75:
        acts += [("1–2w","Nurse/Pharmacist","Med reconciliation; simplify dosing")]
    if has_dep:
        acts += [("1–2w","Clinician","Behavioral activation + activity schedule")]
    if has_scz:
        acts += [("1–4w","Clinician","Relapse plan; early warning signs; caregiver involvement")]
    return acts

bucket = {"High":"High","Moderate–High":"High","Moderate":"Moderate","Low–Moderate":"Low","Low":"Low"}
acts = [ _normalize_action_tuple(a) for a in BASE_ACTIONS[bucket[level]] ]
acts += personalized_actions(X_used.iloc[0], diagnoses, chief)
seen=set(); uniq=[]
for a in acts:
    key=(a[0],a[1],a[2])
    if key not in seen:
        seen.add(key); uniq.append(a)
ORDER={"Today":0,"48h":1,"7d":2,"1–7d":2,"1–2w":3,"2–4w":4,"1–4w":5}
uniq.sort(key=lambda x:(ORDER.get(x[0],99), x[1], x[2]))
st.dataframe(pd.DataFrame(uniq, columns=["Timeline","Owner","Action"]), use_container_width=True)

if level in ["High","Moderate–High"]:
    def make_sop_txt(score, label, actions):
        lines=["Psychiatric Dropout Risk – SOP", f"Risk score: {score}/100 | Risk level: {label}",""]
        for (tl,ow,ac) in actions: lines.append(f"- {tl} | {ow} | {ac}")
        buf=BytesIO("\n".join(lines).encode("utf-8")); buf.seek(0); return buf
    st.download_button("⬇️ Export SOP (TXT)", make_sop_txt(score, level, uniq),
                       file_name="dropout_risk_SOP.txt", mime="text/plain")
    if HAS_DOCX:
        def make_docx(score,label,actions):
            doc=Document(); doc.add_heading('Psychiatric Dropout Risk – SOP', level=1)
            doc.add_paragraph(f"Risk score: {score}/100 | Risk level: {label}")
            t=doc.add_table(rows=1, cols=3); hdr=t.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Timeline','Owner','Action'
            for (tl,ow,ac) in actions:
                r=t.add_row().cells; r[0].text, r[1].text, r[2].text = tl,ow,ac
            buf=BytesIO(); doc.save(buf); buf.seek(0); return buf
        st.download_button("⬇️ Export SOP (Word)", make_docx(score, level, uniq),
                           file_name="dropout_risk_SOP.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ====== What-if 小面板 ======
with st.expander("🧪 What-if: adjust followups/compliance and recompute", expanded=False):
    wf_follow = st.slider("What-if followups (30d count)", 0, 12, int(X_single.at[0,"post_discharge_followups_30d"]))
    wf_comp   = st.slider("What-if compliance", 0.0, 10.0, float(X_single.at[0,"medication_compliance_score"]), 0.5)
    X_wf = X_used.copy()
    X_wf.at[0,"post_discharge_followups_30d"] = wf_follow if use_followups_feature else 0
    X_wf.at[0,"medication_compliance_score"] = wf_comp
    X_wf_al, _ = align_df_to_model(X_wf, model)
    p_m_wf = float(predict_model_proba(X_wf_al)[0])
    p_o_wf, _ = overlay_single_and_drivers(X_wf, base_prob=p_m_wf, include_followup_effect=use_followups_feature)
    p_f_wf = (1.0 - BLEND_W) * p_m_wf + BLEND_W * p_o_wf
    st.write(f"Model={p_m_wf*100:.1f}% | Overlay={p_o_wf*100:.1f}% | Final={p_f_wf*100:.1f}%")

# ====== Batch Prediction ======
st.markdown("---")
st.subheader("Batch Prediction (Excel)")

friendly_cols = [
    "Age","Gender","Diagnoses","Chief Complaint(s)","Bipolar Episode",
    "Length of Stay (days)","Previous Admissions (1y)",
    "Medication Compliance (0–10)","Family Support (0–10)","Financial Strain (0–10)",
    "Post-discharge Followups (30d count)",
    "Living alone","Has case manager","Prior dropout→rehospitalization (1y)",
    "Recent Self-harm","Self-harm During Admission"
]
tpl = pd.DataFrame(columns=friendly_cols)
buf_tpl = BytesIO(); tpl.to_excel(buf_tpl, index=False); buf_tpl.seek(0)
st.download_button("📥 Download Excel Template", buf_tpl, file_name="batch_template_expert.xlsx",
                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

uploaded = st.file_uploader("📂 Upload Excel", type=["xlsx"])
def parse_multi(cell):
    parts = [p.strip() for p in re.split(r"[;,/|]", str(cell)) if p.strip()]
    return parts if parts else []

if uploaded is not None:
    try:
        raw = pd.read_excel(uploaded)
        df = pd.DataFrame(0, index=raw.index, columns=TEMPLATE_COLUMNS, dtype=float)

        def safe(col, default=None):
            return raw[col] if col in raw.columns else default

        # numeric
        map_num = [
            ("Age","age"),
            ("Length of Stay (days)","length_of_stay"),
            ("Previous Admissions (1y)","num_previous_admissions"),
            ("Medication Compliance (0–10)","medication_compliance_score"),
            ("Family Support (0–10)","family_support_score"),
            ("Financial Strain (0–10)","financial_strain_score"),
            ("Post-discharge Followups (30d count)","post_discharge_followups_30d"),
        ]
        for c_in, c_out in map_num:
            if c_in in raw.columns:
                df[c_out] = pd.to_numeric(raw[c_in], errors="coerce")

        # one-hots
        if "Gender" in raw.columns:
            for i, v in raw["Gender"].astype(str).str.strip().items():
                col=f"gender_{v}"; 
                if col in df.columns: df.at[i,col]=1

        if "Diagnoses" in raw.columns:
            for i, cell in raw["Diagnoses"].items():
                for v in parse_multi(cell):
                    col=f"diagnosis_{v}"
                    if col in df.columns: df.at[i,col]=1

        if "Chief Complaint(s)" in raw.columns:
            for i, cell in raw["Chief Complaint(s)"].items():
                for v in parse_multi(cell):
                    col=f"chief_{v}"
                    if col in df.columns: df.at[i,col]=1

        if "Bipolar Episode" in raw.columns:
            for i, v in raw["Bipolar Episode"].astype(str).str.strip().items():
                col=f"bipolar_ep_{v}"
                if col in df.columns: df.at[i,col]=1

        def map_yesno(col_in, prefix):
            if col_in in raw.columns:
                for i, v in raw[col_in].astype(str).str.strip().items():
                    col=f"{prefix}_{v}"
                    if col in df.columns: df.at[i,col]=1

        map_yesno("Living alone","living_alone")
        map_yesno("Has case manager","has_case_manager")
        map_yesno("Prior dropout→rehospitalization (1y)","prior_dropout_rehosp")
        map_yesno("Recent Self-harm","has_recent_self_harm")
        map_yesno("Self-harm During Admission","self_harm_during_admission")

        fill_defaults_batch(df)

        # Pre-planning → followups=0
        if not use_followups_feature:
            df["post_discharge_followups_30d"] = 0

        Xb_al, _ = align_df_to_model(df, model)
        base_probs = predict_model_proba(Xb_al)

        # Overlay vectorized
        def overlay_vec(df_feat: pd.DataFrame, include_followup=True):
            base = _logit_vec(base_probs); lz = base.copy()
            adm = pd.to_numeric(df_feat["num_previous_admissions"], errors="coerce").fillna(DEFAULTS["num_previous_admissions"]).to_numpy()
            comp= pd.to_numeric(df_feat["medication_compliance_score"], errors="coerce").fillna(DEFAULTS["medication_compliance_score"]).to_numpy()
            sup = pd.to_numeric(df_feat["family_support_score"], errors="coerce").fillna(DEFAULTS["family_support_score"]).to_numpy()
            fin = pd.to_numeric(df_feat["financial_strain_score"], errors="coerce").fillna(DEFAULTS["financial_strain_score"]).to_numpy()
            fup = pd.to_numeric(df_feat["post_discharge_followups_30d"], errors="coerce").fillna(DEFAULTS["post_discharge_followups_30d"]).to_numpy()
            los = pd.to_numeric(df_feat["length_of_stay"], errors="coerce").fillna(DEFAULTS["length_of_stay"]).to_numpy()
            agev= pd.to_numeric(df_feat["age"], errors="coerce").fillna(DEFAULTS["age"]).to_numpy()

            lz += POLICY["per_prev_admission"] * np.minimum(adm, 5)
            lz += POLICY["per_point_low_support"] * np.maximum(0.0, 5.0 - sup)
            lz += POLICY["per_point_financial_strain"] * np.maximum(0.0, fin - 5.0)
            lz += POLICY["per_point_low_compliance"] * np.maximum(0.0, 5.0 - comp)
            lz += POLICY["per_point_high_compliance_protect"] * np.maximum(0.0, comp - 7.0)
            if include_followup:
                lz += POLICY["per_followup"] * fup
                lz += POLICY["no_followup_extra"] * (fup == 0)

            lz += np.where(los < 3, POLICY["los_short"],
                     np.where(los <= 21, POLICY["los_mid"],
                     np.where(los <= 28, POLICY["los_mid_high"], POLICY["los_long"])))
            lz += POLICY["age_young"] * (agev < 21) + POLICY["age_old"] * (agev >= 75)

            # social
            for nm, w in [("living_alone_Yes", POLICY["living_alone"]),
                          ("has_case_manager_Yes", POLICY["case_manager_protect"]),
                          ("prior_dropout_rehosp_Yes", POLICY["prior_dropout_rehosp"])]:
                if nm in df_feat.columns:
                    lz += w * (df_feat[nm].to_numpy()==1)

            # Dx
            for dx, w in POLICY["diag"].items():
                col=f"diagnosis_{dx}"
                if col in df_feat.columns: lz += w * (df_feat[col].to_numpy() == 1)

            # Bipolar episode
            for ep, w in POLICY["bip_ep"].items():
                col=f"bipolar_ep_{ep}"
                if col in df_feat.columns: lz += w * (df_feat[col].to_numpy()==1)

            # Chief
            for cf, w in POLICY["chief"].items():
                col=f"chief_{cf}"
                if col in df_feat.columns: lz += w * (df_feat[col].to_numpy()==1)

            sud = (df_feat.get("diagnosis_Substance Use Disorder",0).to_numpy()==1)
            pdm = (df_feat.get("diagnosis_Personality Disorder",0).to_numpy()==1)
            lz += POLICY["x_sud_lowcomp"] * (sud & (comp <= 3))
            lz += POLICY["x_pd_shortlos"] * (pdm & (los < 3))

            delta = np.clip(OVERLAY_SCALE * (lz - base), -DELTA_CLIP, DELTA_CLIP)
            lz2 = base + delta + CAL_LOGIT_SHIFT
            return 1.0 / (1.0 + np.exp(-(lz2 / TEMP)))

        p_overlay_b = overlay_vec(df, include_followup=use_followups_feature)
        p_final_b = (1.0 - BLEND_W) * base_probs + BLEND_W * p_overlay_b

        # self-harm uplift
        hr = df.get("has_recent_self_harm_Yes", 0); ha = df.get("self_harm_during_admission_Yes", 0)
        mask = ((np.array(hr)==1) | (np.array(ha)==1))
        p_final_b[mask] = np.minimum(np.maximum(p_final_b[mask], SOFT_UPLIFT["floor"]) + SOFT_UPLIFT["add"], SOFT_UPLIFT["cap"])

        out = raw.copy()
        out["risk_percent"] = (p_final_b*100).round(1)
        out["risk_score_0_100"] = (p_final_b*100).round().astype(int)

        s = out["risk_score_0_100"].to_numpy()
        levels = np.full(s.shape, "Low", dtype=object)
        levels[s >= 20 - BORDER_BAND] = "Low–Moderate"
        levels[s >= 20 + BORDER_BAND] = "Moderate"
        levels[s >= 40 - BORDER_BAND] = "Moderate–High"
        levels[s >= 40 + BORDER_BAND] = "High"
        out["risk_level"] = levels

        st.dataframe(out, use_container_width=True)
        buf = BytesIO(); out.to_csv(buf, index=False); buf.seek(0)
        st.download_button("⬇️ Download Results (CSV)", buf, "predictions.csv", "text/csv")
    except Exception as e:
        st.error(f"Batch error: {e}")

# ====== Validation（Ablation + Decision curve + Capacity + Fairness）======
st.markdown("---")
st.header("✅ Validation (synthetic hold-out)")

cA, cB = st.columns([2,1])
with cA:
    n_val = st.slider("Number of synthetic patients", 5000, 60000, 20000, 5000)
    seed_val = st.number_input("Random seed", min_value=1, max_value=10**9, value=2024, step=1)
with cB:
    pt_low = st.slider("Decision threshold (Moderate%):", 5, 60, 20, 1)
    pt_high = st.slider("Decision threshold (High%):", 20, 90, 40, 1)
run_val = st.button("Run validation")

def generate_synth_holdout(n=20000, seed=2024):
    rng = np.random.default_rng(seed)
    df = pd.DataFrame(0, index=range(n), columns=TEMPLATE_COLUMNS, dtype=float)
    df["age"] = rng.integers(16, 85, n)
    base_los = rng.normal(22.0, 8.0, n)
    base_los = np.clip(base_los, 1, 60)
    df["length_of_stay"] = base_los
    df["num_previous_admissions"] = rng.poisson(1.0, n).clip(0, 12)
    df["medication_compliance_score"] = rng.normal(6.0, 2.5, n).clip(0, 10)
    df["family_support_score"] = rng.normal(5.0, 2.5, n).clip(0, 10)
    df["financial_strain_score"] = rng.normal(5.0, 2.5, n).clip(0, 10)
    df["post_discharge_followups_30d"] = rng.integers(0, 6, n)

    # gender
    idx_gender = rng.integers(0, len(GENDER_LIST), n)
    for i, g in enumerate(GENDER_LIST): df.loc[idx_gender == i, f"gender_{g}"] = 1

    # Dx
    idx_primary = rng.integers(0, len(DIAG_LIST), n)
    for i, d in enumerate(DIAG_LIST): df.loc[idx_primary == i, f"diagnosis_{d}"] = 1
    for d, pr in {"Substance Use Disorder": 0.22, "Depression": 0.27, "Anxiety": 0.22, "PTSD": 0.12}.items():
        df.loc[rng.random(n) < pr, f"diagnosis_{d}"] = 1

    # LOS shift by Dx
    los = df["length_of_stay"].to_numpy()
    los += 5.0 * (df["diagnosis_Schizophrenia"]==1).to_numpy()
    los += 3.0 * (df["diagnosis_Bipolar"]==1).to_numpy()
    df["length_of_stay"] = np.clip(los, 1, 90)

    # Chief
    for c in CHIEF_LIST:
        df.loc[rng.random(n) < 0.15, f"chief_{c}"] = 1
    chief_cols = [f"chief_{c}" for c in CHIEF_LIST]
    none_chief = (df[chief_cols].sum(axis=1) == 0)
    df.loc[none_chief, "chief_Other/Unknown"] = 1

    # Bipolar episode
    probs_bip = [0.70, 0.07, 0.10, 0.05, 0.08]
    choices = rng.choice(len(BIPOLAR_EP_LIST), size=n, p=probs_bip)
    for i, b in enumerate(BIPOLAR_EP_LIST):
        df.loc[choices == i, f"bipolar_ep_{b}"] = 1

    # social
    df["living_alone_Yes"] = (rng.random(n) < 0.25).astype(int); df["living_alone_No"] = 1 - df["living_alone_Yes"]
    df["has_case_manager_Yes"] = (rng.random(n) < 0.35).astype(int); df["has_case_manager_No"] = 1 - df["has_case_manager_Yes"]
    df["prior_dropout_rehosp_Yes"] = (rng.random(n) < 0.12).astype(int); df["prior_dropout_rehosp_No"] = 1 - df["prior_dropout_rehosp_Yes"]

    # self-harm
    r1, r2 = rng.integers(0, 2, n), rng.integers(0, 2, n)
    df.loc[r1 == 1, "has_recent_self_harm_Yes"] = 1; df.loc[r1 == 0, "has_recent_self_harm_No"] = 1
    df.loc[r2 == 1, "self_harm_during_admission_Yes"] = 1; df.loc[r2 == 0, "self_harm_during_admission_No"] = 1

    fill_defaults_batch(df)
    # ground truth (same logic as training)
    beta0 = -0.50
    prev_ge2 = (pd.to_numeric(df["num_previous_admissions"], errors="coerce").fillna(0).to_numpy() >= 2).astype(np.float32)

    def col(df_, name): return pd.to_numeric(df_[name], errors="coerce").fillna(0).to_numpy().astype(np.float32)

    logit = (beta0
        + 0.80*col(df,"has_recent_self_harm_Yes")
        + 0.60*col(df,"self_harm_during_admission_Yes")
        + 0.60*prev_ge2
        - 0.25*col(df,"medication_compliance_score")
        - 0.20*col(df,"family_support_score")
        + 0.12*col(df,"financial_strain_score")
        - 0.15*col(df,"post_discharge_followups_30d")
        + 0.05*col(df,"length_of_stay"))

    for dx, w in POLICY["diag"].items():
        nm=f"diagnosis_{dx}"
        if nm in df.columns: logit += w*col(df,nm)

    for ep, w in POLICY["bip_ep"].items():
        nm=f"bipolar_ep_{ep}"
        if nm in df.columns: logit += w*col(df,nm)

    for cf, w in POLICY["chief"].items():
        nm=f"chief_{cf}"
        if nm in df.columns: logit += w*col(df,nm)

    logit += 0.20*col(df,"living_alone_Yes") - 0.25*col(df,"has_case_manager_Yes") + 0.40*col(df,"prior_dropout_rehosp_Yes")

    if "diagnosis_Substance Use Disorder" in df.columns:
        logit += 0.30*((col(df,"diagnosis_Substance Use Disorder")==1) & (col(df,"medication_compliance_score")<=3))
    if "diagnosis_Personality Disorder" in df.columns:
        logit += 0.10*((col(df,"diagnosis_Personality Disorder")==1) & (col(df,"length_of_stay")<3))

    noise = np.random.default_rng(seed+1).normal(0.0, 0.35, n).astype(np.float32)
    p_true = 1.0 / (1.0 + np.exp(-(logit + noise)))
    y_true = (np.random.default_rng(seed+2).random(n) < p_true).astype(int)

    return df, y_true

def plot_roc_pr(y, p_list, labels):
    from sklearn.metrics import roc_curve, auc, precision_recall_curve, average_precision_score
    fig1, ax1 = plt.subplots()
    for p,l in zip(p_list, labels):
        fpr, tpr, _ = roc_curve(y, p); roc = auc(fpr, tpr)
        ax1.plot(fpr, tpr, label=f"{l} AUC={roc:.3f}")
    ax1.plot([0,1],[0,1],"--")
    ax1.set_xlabel("FPR"); ax1.set_ylabel("TPR"); ax1.set_title("ROC")
    ax1.legend(loc="lower right"); st.pyplot(fig1, clear_figure=True)

    fig2, ax2 = plt.subplots()
    for p,l in zip(p_list, labels):
        prec, rec, _ = precision_recall_curve(y, p); ap = average_precision_score(y, p)
        ax2.plot(rec, prec, label=f"{l} AP={ap:.3f}")
    ax2.set_xlabel("Recall"); ax2.set_ylabel("Precision"); ax2.set_title("PR")
    ax2.legend(loc="upper right"); st.pyplot(fig2, clear_figure=True)

def ece(y, p, n_bins=10):
    bins = np.linspace(0.0,1.0,n_bins+1)
    idx = np.digitize(p, bins)-1
    err=0.0
    for b in range(n_bins):
        m=(idx==b)
        if m.sum()==0: continue
        fp=p[m].mean(); tp=y[m].mean()
        err += m.mean()*abs(tp-fp)
    return float(err)

def confusion(y, p, thr):
    from sklearn.metrics import confusion_matrix
    yhat = (p>=thr).astype(int)
    tn,fp,fn,tp = confusion_matrix(y,yhat).ravel()
    return tn,fp,fn,tp

def decision_curve(y, p):
    ths = np.linspace(0.05,0.60,56)
    N=len(y)
    nb=[]
    for t in ths:
        yhat = (p>=t).astype(int)
        tp = ((yhat==1)&(y==1)).sum(); fp=((yhat==1)&(y==0)).sum()
        nb.append((tp/N) - (fp/N)*(t/(1-t)))
    return ths, np.array(nb)

if run_val:
    try:
        from sklearn.metrics import roc_auc_score, average_precision_score, brier_score_loss
    except Exception as e:
        st.error(f"Need scikit-learn: {e}")
    else:
        with st.spinner("Generating data & evaluating..."):
            df_syn, y_true = generate_synth_holdout(n_val, seed_val)
            if not use_followups_feature:
                df_syn["post_discharge_followups_30d"] = 0
            Xa, _ = align_df_to_model(df_syn, model)
            p_model_v = predict_model_proba(Xa)

            def overlay_only_vec(df_feat, base_probs):
                base = _logit_vec(base_probs); lz = base.copy()
                adm = pd.to_numeric(df_feat["num_previous_admissions"], errors="coerce").fillna(DEFAULTS["num_previous_admissions"]).to_numpy()
                comp= pd.to_numeric(df_feat["medication_compliance_score"], errors="coerce").fillna(DEFAULTS["medication_compliance_score"]).to_numpy()
                sup = pd.to_numeric(df_feat["family_support_score"], errors="coerce").fillna(DEFAULTS["family_support_score"]).to_numpy()
                fin = pd.to_numeric(df_feat["financial_strain_score"], errors="coerce").fillna(DEFAULTS["financial_strain_score"]).to_numpy()
                fup = pd.to_numeric(df_feat["post_discharge_followups_30d"], errors="coerce").fillna(DEFAULTS["post_discharge_followups_30d"]).to_numpy()
                los = pd.to_numeric(df_feat["length_of_stay"], errors="coerce").fillna(DEFAULTS["length_of_stay"]).to_numpy()
                agev= pd.to_numeric(df_feat["age"], errors="coerce").fillna(DEFAULTS["age"]).to_numpy()

                lz += POLICY["per_prev_admission"] * np.minimum(adm, 5)
                lz += POLICY["per_point_low_support"] * np.maximum(0.0, 5.0 - sup)
                lz += POLICY["per_point_financial_strain"] * np.maximum(0.0, fin - 5.0)
                lz += POLICY["per_point_low_compliance"] * np.maximum(0.0, 5.0 - comp)
                lz += POLICY["per_point_high_compliance_protect"] * np.maximum(0.0, comp - 7.0)
                lz += POLICY["per_followup"] * fup
                lz += POLICY["no_followup_extra"] * (fup == 0)

                lz += np.where(los < 3, POLICY["los_short"],
                        np.where(los <= 21, POLICY["los_mid"],
                        np.where(los <= 28, POLICY["los_mid_high"], POLICY["los_long"])))
                lz += POLICY["age_young"] * (agev < 21) + POLICY["age_old"]*(agev >= 75)

                for nm, w in [("living_alone_Yes", POLICY["living_alone"]),
                              ("has_case_manager_Yes", POLICY["case_manager_protect"]),
                              ("prior_dropout_rehosp_Yes", POLICY["prior_dropout_rehosp"])]:
                    if nm in df_feat.columns:
                        lz += w * (df_feat[nm].to_numpy()==1)

                for dx,w in POLICY["diag"].items():
                    col=f"diagnosis_{dx}"
                    if col in df_feat.columns: lz += w * (df_feat[col].to_numpy()==1)

                for ep,w in POLICY["bip_ep"].items():
                    col=f"bipolar_ep_{ep}"
                    if col in df_feat.columns: lz += w * (df_feat[col].to_numpy()==1)

                for cf,w in POLICY["chief"].items():
                    col=f"chief_{cf}"
                    if col in df_feat.columns: lz += w * (df_feat[col].to_numpy()==1)

                sud = (df_feat.get("diagnosis_Substance Use Disorder",0).to_numpy()==1)
                pdm = (df_feat.get("diagnosis_Personality Disorder",0).to_numpy()==1)
                lz += POLICY["x_sud_lowcomp"] * (sud & (comp <= 3))
                lz += POLICY["x_pd_shortlos"] * (pdm & (los < 3))

                delta = np.clip(OVERLAY_SCALE * (lz - base), -DELTA_CLIP, DELTA_CLIP)
                lz2 = base + delta + CAL_LOGIT_SHIFT
                return 1.0 / (1.0 + np.exp(-(lz2 / TEMP)))

            p_overlay_v = overlay_only_vec(df_syn, p_model_v)
            p_final_v = (1.0 - BLEND_W) * p_model_v + BLEND_W * p_overlay_v

            auc_m = roc_auc_score(y_true, p_model_v); auc_o = roc_auc_score(y_true, p_overlay_v); auc_f = roc_auc_score(y_true, p_final_v)
            ap_m  = average_precision_score(y_true, p_model_v); ap_o  = average_precision_score(y_true, p_overlay_v); ap_f  = average_precision_score(y_true, p_final_v)
            br_m  = brier_score_loss(y_true, p_model_v); br_o  = brier_score_loss(y_true, p_overlay_v); br_f  = brier_score_loss(y_true, p_final_v)
            ece_m = ece(y_true, p_model_v); ece_o = ece(y_true, p_overlay_v); ece_f = ece(y_true, p_final_v)
            st.subheader("Ablation metrics")
            st.dataframe(pd.DataFrame([
                {"Model":"Model","AUC":auc_m,"PR-AUC":ap_m,"Brier":br_m,"ECE(10bins)":ece_m},
                {"Model":"Overlay-only","AUC":auc_o,"PR-AUC":ap_o,"Brier":br_o,"ECE(10bins)":ece_o},
                {"Model":"Blend (Final)","AUC":auc_f,"PR-AUC":ap_f,"Brier":br_f,"ECE(10bins)":ece_f},
            ]).round(3), use_container_width=True)

            st.subheader("Curves")
            plot_roc_pr(y_true, [p_model_v, p_overlay_v, p_final_v], ["Model","Overlay","Final"])

            ths_m, nb_m = decision_curve(y_true, p_model_v)
            ths_o, nb_o = decision_curve(y_true, p_overlay_v)
            ths_f, nb_f = decision_curve(y_true, p_final_v)
            fig, ax = plt.subplots()
            ax.plot(ths_m, nb_m, label="Model")
            ax.plot(ths_o, nb_o, label="Overlay")
            ax.plot(ths_f, nb_f, label="Final")
            ax.axhline(0, ls="--"); ax.set_xlabel("Threshold"); ax.set_ylabel("Net benefit"); ax.set_title("Decision Curve")
            ax.legend(); st.pyplot(fig, clear_figure=True)

            thr_mod = pt_low/100.0; thr_hi = pt_high/100.0
            st.subheader("Operational (binary @ thresholds)")
            c1, c2 = st.columns(2)
            with c1:
                st.markdown(f"**Moderate ≥{pt_low}%**")
                for name, p in [("Model",p_model_v),("Overlay",p_overlay_v),("Final",p_final_v)]:
                    tn,fp,fn,tp = confusion(y_true, p, thr_mod)
                    st.write(f"{name} — TN:{tn} FP:{fp} FN:{fn} TP:{tp}")
            with c2:
                st.markdown(f"**High ≥{pt_high}%**")
                for name, p in [("Model",p_model_v),("Overlay",p_overlay_v),("Final",p_final_v)]:
                    tn,fp,fn,tp = confusion(y_true, p, thr_hi)
                    st.write(f"{name} — TN:{tn} FP:{fp} FN:{fn} TP:{tp}")

            # Capacity & time load
            st.subheader("Capacity / Time load (per 1,000 patients)")
            N = 1000
            def count_at(p, thr): return int(((p>=thr).sum() / len(p)) * N)
            mod_cnt = count_at(p_final_v, thr_mod); high_cnt = count_at(p_final_v, thr_hi)
            st.caption("Assumptions (editable):")
            t_outreach = st.number_input("Nurse outreach (min)", 5, 60, 15, 5)
            t_sched = st.number_input("Scheduler booking (min)", 2, 30, 5, 1)
            t_pharm = st.number_input("Pharmacist review (min)", 5, 60, 20, 5)
            hours = (mod_cnt*(t_outreach+t_sched) + (high_cnt)*(t_pharm)) / 60.0
            st.write(f"Flagged Moderate+ per 1,000: **{mod_cnt}** ; High: **{high_cnt}** → ~ **{hours:.1f} hours** total effort.")

            # Fairness：性別/年齡段/主診斷
            st.subheader("Fairness (AUC / ECE by subgroup, Final)")
            def subgroup_auc_ece(df_feat, y, p, mask):
                if mask.sum()<100: return np.nan, np.nan
                from sklearn.metrics import roc_auc_score
                return float(roc_auc_score(y[mask], p[mask])), float(ece(y[mask], p[mask], 10))
            agev = pd.to_numeric(df_syn["age"], errors="coerce").fillna(40).to_numpy()
            bands = {"<30": (agev<30), "30–59": ((agev>=30)&(agev<60)), "≥60": (agev>=60)}
            rows=[]
            for g in GENDER_LIST:
                mask = (df_syn.get(f"gender_{g}",0).to_numpy()==1)
                a,e = subgroup_auc_ece(df_syn, y_true, p_final_v, mask)
                rows.append({"group":"Gender", "value":g, "AUC":a, "ECE":e})
            for k,m in bands.items():
                a,e = subgroup_auc_ece(df_syn, y_true, p_final_v, m)
                rows.append({"group":"AgeBand", "value":k, "AUC":a, "ECE":e})
            diag_cols=[f"diagnosis_{d}" for d in DIAG_LIST]
            prim = np.argmax(df_syn[diag_cols].to_numpy(), axis=1)
            for i,d in enumerate(DIAG_LIST):
                mask = (prim==i)
                a,e = subgroup_auc_ece(df_syn, y_true, p_final_v, mask)
                rows.append({"group":"PrimaryDx", "value":d, "AUC":a, "ECE":e})
            st.dataframe(pd.DataFrame(rows).round(3), use_container_width=True)

        st.success("Validation finished.")

# ====== Vignettes（for expert review）======
st.markdown("---")
st.header("🧾 Vignettes template (for expert review)")
def _mk_vignette_row(age, gender, diags, chief_list, bip, los, prev, comp, rsh, shadm, sup, fin, fup, alone, cm, prior):
    return {
        "Age": age, "Gender": gender, "Diagnoses": ", ".join(diags),
        "Chief Complaint(s)": ", ".join(chief_list), "Bipolar Episode": bip,
        "Length of Stay (days)": los, "Previous Admissions (1y)": prev,
        "Medication Compliance (0–10)": comp,
        "Family Support (0–10)": sup,
        "Financial Strain (0–10)": fin,
        "Post-discharge Followups (30d count)": fup,
        "Living alone": alone, "Has case manager": cm, "Prior dropout→rehospitalization (1y)": prior,
        "Recent Self-harm": rsh, "Self-harm During Admission": shadm,
        "Expert Risk (Low/Moderate/High or 0–100)": ""
    }
def build_vignettes_df(n=20, seed=77):
    rng = np.random.default_rng(seed); base=[]
    protos = [
        (19,"Female",["Depression"],["Suicidal ideation/attempt"],"N/A",14,0,3,"Yes","No",2,6,0,"No","No","No"),
        (28,"Male",["Substance Use Disorder"],["Self-harm (non-suicidal)"],"N/A",10,3,2,"No","No",3,7,0,"Yes","No","Yes"),
        (35,"Male",["Bipolar"],["Severe agitation/mania"],"Manic",24,1,4,"No","No",5,5,1,"No","No","No"),
        (42,"Female",["Personality Disorder"],["Self-harm (non-suicidal)"],"N/A",8,2,3,"Yes","No",4,5,0,"No","No","No"),
        (55,"Male",["Schizophrenia"],["Psychosis/diagnostic workup"],"N/A",28,4,5,"No","No",5,4,2,"No","Yes","Yes"),
        (63,"Female",["PTSD"],["Suicidal ideation/attempt"],"N/A",18,1,6,"Yes","No",6,3,1,"No","No","No"),
    ]
    for i in range(len(protos)):
        base.append(_mk_vignette_row(*protos[i]))
    while len(base) < n:
        age = int(np.clip(rng.normal(40,15), 18, 90))
        gender = GENDER_LIST[int(rng.integers(0,len(GENDER_LIST)))]
        k = int(rng.integers(1,3)); diags = list(rng.choice(DIAG_LIST, size=k, replace=False))
        k2 = int(rng.integers(1,3)); chiefs = list(rng.choice(CHIEF_LIST, size=k2, replace=False))
        bip = rng.choice(BIPOLAR_EP_LIST)
        los = int(np.clip(rng.normal(22,8),1,60)); prev = int(np.clip(rng.poisson(1.0),0,8))
        comp = float(np.clip(rng.normal(6,2.5),0,10)); rsh = rng.choice(["Yes","No"]); shadm = rng.choice(["Yes","No"])
        sup = float(np.clip(rng.normal(5,2.5),0,10)); fin = float(np.clip(rng.normal(5,2.5),0,10))
        fup = int(np.clip(rng.integers(0,6),0,12)); alone = rng.choice(["Yes","No"]); cm = rng.choice(["Yes","No"]); prior=rng.choice(["Yes","No"])
        base.append(_mk_vignette_row(age, gender, diags, chiefs, bip, los, prev, comp, rsh, shadm, sup, fin, fup, alone, cm, prior))
    return pd.DataFrame(base[:n])

vdf = build_vignettes_df(20, 77)
buf_v = BytesIO(); vdf.to_excel(buf_v, index=False); buf_v.seek(0)
st.download_button("📥 Download Vignettes (20 cases, Excel)", buf_v,
                   file_name="vignettes_20_expert.xlsx",
                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ====== Data dictionary ======
# ====== Data dictionary ======
with st.expander("📚 Data dictionary / Definitions", expanded=False):
    st.markdown(f"""
- **Medication Compliance (0–10)**：0=幾乎不服藥；10=幾乎完全依從（近 1 個月）
- **Family Support (0–10)**：0=非常不足；10=非常充足
- **Financial Strain (0–10)**：0=無壓力；10=極高壓力
- **{FOLLOWUPS_LABEL}**：出院後 30 天內的門診/電話/社工接觸次數（本系統一律以 30 天統一定義）
- **Self-harm flags**：最近自傷 / 住院期間自傷
- **Chief Complaint(s)**：本次住院的主要問題（可複選），如自殺意念/自傷、攻擊行為、躁動等
- **Bipolar Episode**：Manic / Depressive / Mixed / Hypomanic / N/A
- **Pre-planning 模式**：為避免洩漏，計算時忽略 30 天追蹤次數特徵
- **Final Probability**：Model 與 Policy Overlay 混合（可調 BLEND），並含必要的 safety uplift（自傷相關個案至少 0.6 起跳後再疊加）
- **Risk bins**：Low(<~13)、Low–Moderate(≈13–27)、Moderate(≈27–33)、Moderate–High(≈33–47)、High(>~47)；實際由門檻與緩衝帶計算（可在 Validation 區調整）
- **Decision Curve (Net Benefit)**：NB = TP/N − FP/N × (t/(1−t))，比較模型決策與 treat-all / treat-none 的淨效益
- **ECE (Expected Calibration Error)**：分箱後 |預測機率 − 實際發生率| 的加權平均，衡量校準度
- **用途定位**：臨床決策輔助（CDSS），非取代臨床判斷；高風險個案需人工覆核並保留稽核紀錄
""")

st.caption("Demo 版以合成資料 + 臨床啟發規則做示範；臨床部署前需以院內實證數據訓練/驗證並通過 IRB 與資安稽核。")
